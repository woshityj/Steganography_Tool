from docx import Document
from docx.shared import RGBColor
from steganograpy_utility import to_bin


def encode(path, msg, bit):
    docu = Document(path)
    msg += "====="
    binary_msg = to_bin(msg)
    length = len(binary_msg)
    # check size
    max_length = 0
    for para in docu.paragraphs:
        max_length += len(para.text)
    if length > max_length:
        raise ValueError(f"Message length ({length}) exceeded {max_length}, please adjust the message or bit.")

    index = 0

    for para in docu.paragraphs:
        if index >= length:
            break
        runs = para.runs
        para.clear()
        for run in runs:
            text = run.text
            if run.font.color.rgb is not None:
                color = to_bin(run.font.color.rgb)
            else:
                color = ['00000000', '00000000', '00000000']
            for _, char in enumerate(text):
                if index >= length:
                    para.add_run(char)
                    continue
                run = para.add_run(char)
                for i in range(3):
                    data = ""
                    for b in range(bit):
                        if index < length:
                            data += binary_msg[index]
                            index += 1
                        else:
                            data += color[i][-bit + b]
                    color[i] = color[i][:-bit] + data
                run.font.color.rgb = RGBColor(int(color[0], 2), int(color[1], 2), int(color[2], 2))
    return docu


def decode(path, bit):
    docu = Document(path)
    binary_msg = ""
    msg = ""
    # for each run
    for para in docu.paragraphs:
        for run in para.runs:
            if run.font.color.rgb is None:
                continue
            for c in run.font.color.rgb:
                binary_color = to_bin(c)
                binary_msg += binary_color[-bit:]
                if len(binary_msg) >= 8:
                    msg += chr(int(binary_msg[:8], 2))
                    if msg.endswith("====="):
                        return msg[:-5]
                    binary_msg = binary_msg[8:]

def save(path, docu):
    if docu:
        docu.save(path)


save("testingfiles/out.docx", encode('testingfiles/SampleWord.docx', """hello world!!
hello world!!
hello world!!
hello world!!
hello world!!
""", 5))
print(decode("testingfiles/out.docx", 5))
