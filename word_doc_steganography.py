from docx.oxml import OxmlElement

def setHiddenProperty(p):
    pPr = OxmlElement('w:pPr') # paragraph property
    rPr = OxmlElement('w:rPr') # run property
    v = OxmlElement('w:vanish') # hidden
    rPr.append(v)
    pPr.append(rPr)
    p._p.append(pPr)

def hiddenParagraphTest(filepath, text):
    import docx
    d = docx.Document(filepath)
    #p = d.add_paragraph("Visible")

    p = d.add_paragraph()
    setHiddenProperty(p)  # set paragraph hidden property
    r = p.add_run()
    r.text = text
    r.font.hidden = False

    return d

    #p = d.add_paragraph("Visible")
    d.save(filepath)

def findParagraph(filepath):
    import docx
    d = docx.Document(filepath)
    last_p_in_document = d.paragraphs[-1]
    return last_p_in_document.text