import os
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import messagebox
from tkinterdnd2 import DND_FILES, TkinterDnD
import customtkinter
from tkinter.filedialog import asksaveasfilename, askopenfilename
import wave
import vlc
from tkVideoPlayer import TkinterVideo
import docx

import cv2
import image_steganography as ims
import document_steganography as dms
import audio_steganography as auds
import gif_steganography as gis
import video_steganography as vids
import word_doc_steganography as wds
import word_steganography as wd_lsb_s
import xlsx_steganography as xls
import encrypt as enc

supported_types = {
    '.png': (ims.png_encode, ims.png_decode),
    '.gif': (gis.gif_encode, gis.gif_decode),
    '.bmp': (ims.bmp_encode, ims.bmp_decode),
    '.wav': (auds.encode_aud_data, auds.decode_aud_data),
    '.mp3': (auds.encode_aud_data, auds.decode_aud_data),
    '.mp4': (vids.encode_video, vids.decode_video),
    '.txt': (dms.decode),
    '.xlsx': (xls.encode, xls.decode),
    '.docx': (wd_lsb_s.encode, wd_lsb_s.decode)
}


def get_path(filetype):
    file_path = askopenfilename(
        title='Open a file',
        initialdir='/',
        filetypes=filetype
    )
    if file_path:
        return file_path


def get_drop(event, filetype):
    file_path = event.data
    file_path = file_path.removeprefix('{').removesuffix('}')
    if os.path.isfile(file_path):
        _, ext = os.path.splitext(file_path)
    else:
        messagebox.showerror("Error", f"Cannot open file from {file_path}")
        return
    if ext.lower() not in list(filetype):
        messagebox.showerror("Error", f"File type {ext} not supported.")
        return
    return file_path


class MainMenu(customtkinter.CTkFrame):
    def __init__(self, master, controller, **kwargs):
        super().__init__(master, **kwargs)
        self.grid_columnconfigure(0, weight=1)

        # add widgets onto the frame, for example:
        self.label = customtkinter.CTkLabel(self, text="Steganography",
                                            font=customtkinter.CTkFont(size=20, weight="bold"))
        self.label.grid(row=0, column=0, padx=20, pady=10)
        self.img_steg_button = customtkinter.CTkButton(self, text="Image Steganography",
                                                       command=lambda: controller.show_frame(ImagePage))
        self.img_steg_button.grid(row=1, column=0, padx=20, pady=5, sticky="ew")

        self.doc_steg_button = customtkinter.CTkButton(self, text="Document Steganography",
                                                       command=lambda: controller.show_frame(DocumentPage))
        self.doc_steg_button.grid(row=2, column=0, padx=20, pady=5, sticky="ew")

        self.img_steg_button = customtkinter.CTkButton(self, text="Audio Steganography",
                                                       command=lambda: controller.show_frame(AudioPage))
        self.img_steg_button.grid(row=3, column=0, padx=20, pady=5, sticky="ew")

        self.img_steg_button = customtkinter.CTkButton(self, text="Video Steganography",
                                                       command=lambda: controller.show_frame(VideoPage))
        self.img_steg_button.grid(row=4, column=0, padx=20, pady=5, sticky="ew")


def open_img(name, img):
    cv2.imshow(name, img)


def play_gif(fn, frames, label, duration, imagepage):
    f = frames[fn]
    fn += 1
    if fn == len(frames):
        fn = 0
    label.configure(image=f)
    return imagepage.after(duration, lambda: play_gif(fn, frames, label, duration, imagepage))


def clear_text(textarea):
    textarea.delete('1.0', tk.END)


class GifWindow(tk.Toplevel):
    def __init__(self, master, path):
        super().__init__(master=master)
        self.title(path)
        gif = Image.open(path)
        self.duration = gif.info['duration']
        self.frames = [tk.PhotoImage(file=path, format='gif -index %i' % i) for i in range(gif.n_frames)]
        self.label = tk.Label(self, image=self.frames[0], text="")
        self.label.pack()
        self.after(0, lambda: play_gif(0, self.frames, self.label, self.duration, self))


class ImagePage(customtkinter.CTkFrame):
    def __init__(self, master, controller, **kwargs):
        super().__init__(master, **kwargs)

        self.coverPath = None
        self.encoded_image_label = customtkinter.CTkLabel(self)
        self.image_label = customtkinter.CTkLabel(self)
        self.gif_label = tk.Label(self)
        self.encoded_gif_label = tk.Label(self)
        self.gif_after_id = None
        self.encoded_gif_after_id = None

        self.grid_columnconfigure((0, 1), weight=1)

        self.back_button = customtkinter.CTkButton(self, text="Back", command=lambda: controller.show_frame(MainMenu))
        self.back_button.grid(row=0, column=0, columnspan=2, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Image Steganography",
                                            font=customtkinter.CTkFont(size=20, weight="bold"))
        self.label.grid(row=1, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Settings", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=2, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Set Number of Bits", font=customtkinter.CTkFont(size=15))
        self.label.grid(row=3, column=0, padx=20, pady=2)
        self.bits_option_menu = customtkinter.CTkOptionMenu(self, values=["1", "2", "3", "4", "5"])
        self.bits_option_menu.grid(row=4, column=0, padx=20, pady=(10, 10))

        self.checkbox = customtkinter.CTkCheckBox(self, text="Save Text")
        self.checkbox.grid(row=4, column=1, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Cover", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=5, column=0, padx=20, pady=10, columnspan=2)
        self.upload_button = customtkinter.CTkButton(self, text="Click here or Drop an Image here to upload an Image",
                                                     command=self.cover_on_change)
        self.upload_button.grid(row=6, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_button.drop_target_register(DND_FILES)
        self.upload_button.dnd_bind("<<Drop>>", self.cover_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Secret Message", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=7, column=0, padx=20, pady=10, columnspan=2)
        self.secret_message = customtkinter.CTkTextbox(self, height=100, width=500)
        self.secret_message.grid(row=8, column=0, sticky='w', pady=10, columnspan=2)

        self.upload_txt_button = customtkinter.CTkButton(self,
                                                         text="Click here or Drop an txt here to upload a secret message",
                                                         command=self.txt_on_change)
        self.upload_txt_button.grid(row=9, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_txt_button.drop_target_register(DND_FILES)
        self.upload_txt_button.dnd_bind("<<Drop>>", self.txt_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Key", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=10, column=0, padx=20, pady=10, columnspan=2)
        self.key_message = customtkinter.CTkTextbox(self, height=40, width=500)
        self.key_message.grid(row=11, column=0, sticky='w', pady=10, columnspan=2)

        self.encode_button = customtkinter.CTkButton(self, text="Encode Image", command=self.encode_image)
        self.encode_button.grid(row=12, column=0, padx=20, pady=10)

        self.decode_button = customtkinter.CTkButton(self, text="Decode Image", command=self.decode_image)
        self.decode_button.grid(row=12, column=1, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Uploaded Image", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=13, column=0, padx=20, pady=10, columnspan=1)

        self.label = customtkinter.CTkLabel(self, text="Encoded Image", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=13, column=1, padx=20, pady=10, columnspan=1)

    def txt_on_drop(self, event):
        self.txtPath = get_drop(event, ('', '.txt'))
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def txt_on_change(self):
        self.txtPath = get_path([('', '*.txt')])
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def cover_on_drop(self, event):
        self.coverPath = get_drop(event, ('.png', '.bmp', '.gif'))
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.show(self.coverPath)

    def cover_on_change(self):
        self.coverPath = get_path([('', '*.png'), ('', '*.bmp'), ('', '*.gif')])
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.show(self.coverPath)

    def show(self, coverPath):
        if self.gif_after_id is not None:
            self.after_cancel(self.gif_after_id)
            self.gif_after_id = None
        self.image_label.grid_forget()
        self.gif_label.grid_forget()
        if coverPath.lower().endswith(('.png', '.bmp')):
            self.uploaded_image = customtkinter.CTkImage(Image.open(coverPath), size=(100, 100))
            self.image_label = customtkinter.CTkLabel(self, image=self.uploaded_image, text="")
            self.image_label.bind("<Double-Button-1>", lambda event: open_img(coverPath, cv2.imread(coverPath)))
            self.image_label.grid(row=14, column=0)
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=15, column=0)

        if coverPath.lower().endswith('.gif'):
            self.gif = Image.open(coverPath)
            self.frames = [tk.PhotoImage(file=coverPath, format='gif -index %i' % i) for i in range(self.gif.n_frames)]
            self.gif_label = tk.Label(self, image=self.frames[0], text="", height=100, width=100)
            self.gif_label.bind("<Double-Button-1>", lambda e: GifWindow(self, coverPath))
            self.gif_label.grid(row=14, column=0)
            self.gif_after_id = self.after(0,
                                           lambda: play_gif(0, self.frames, self.gif_label, self.gif.info['duration'],
                                                            self))

    def encode_image(self):
        # Error Handling if the User has not select a Cover File
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return

        _, ext = os.path.splitext(self.coverPath)

        if len(self.secret_message.get('1.0', 'end-1c')) == 0:
            messagebox.showerror("Error", "Please enter a secret message.")
            return

        self.payloadText = self.secret_message.get('1.0', 'end-1c')
        message = self.key_message.get('1.0', 'end-1c')
        if (len(message) > 0):
            if (enc.has_repeating_characters(message)):
                messagebox.showerror("Error",
                                     "Secret key has repeating characters, please use a unique string sequence")
                return
            else:
                self.payloadText = enc.encryptMessage(self.payloadText, message)
        try:
            if ext.lower() == '.png' or ext.lower() == '.gif' or ext.lower() == '.bmp':
                self.encoded = supported_types[ext.lower()][0](self.coverPath, self.payloadText,
                                                               int(self.bits_option_menu.get()))

        # Error Handling if the payload is too large to keep in the Cover File
        except ValueError:
            messagebox.showerror("Error", "Insufficient bytes, need a bigger image or reduce payload")
            return
        self.save_as(ext)

    def decode_image(self):
        # Error Handling if the User has not select a Cover File
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return

        _, ext = os.path.splitext(self.coverPath)
        if ext.lower() == '.png' or ext.lower() == '.gif' or ext.lower() == '.bmp':
            message = self.key_message.get('1.0', 'end-1c')
            text = supported_types[ext.lower()][1](self.coverPath, int(self.bits_option_menu.get()))
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                else:
                    text = enc.decryptMessage(text, message)
            if text == 0:
                messagebox.showerror("Error", "Wrong Key")
                return
            clear_text(self.secret_message)
        self.secret_message.insert('1.0', text)
        self.save_text()

        # if self.encoded_image_label is not None:
        #     self.encoded_image_label.destroy()

    def save_as(self, ext):
        # save as prompt
        filename = asksaveasfilename(defaultextension=ext, filetypes=[("Same as original", ext), ("All Files", "*.*")])
        self.encoded_image_label.grid_forget()
        self.encoded_gif_label.grid_forget()
        if self.encoded_gif_after_id is not None:
            self.after_cancel(self.encoded_gif_after_id)
            self.encoded_gif_after_id = None
        if filename:
            if filename.endswith('.png'):
                # save image
                cv2.imwrite(filename, self.encoded)
                # show image as per requested in spec
                self.encoded_image = customtkinter.CTkImage(Image.open(filename), size=(100, 100))
                self.encoded_image_label = customtkinter.CTkLabel(self, image=self.encoded_image, text="")
                self.encoded_image_label.bind("<Double-Button-1>", lambda event: open_img(filename, self.encoded))
                self.encoded_image_label.grid(row=14, column=1)
            elif filename.endswith('.bmp'):
                # save image
                with open(filename, "wb") as image:
                    image.write(self.encoded)
                # show image as per requested in spec
                self.encoded_image = customtkinter.CTkImage(Image.open(filename), size=(100, 100))
                self.encoded_image_label = customtkinter.CTkLabel(self, image=self.encoded_image, text="")
                self.encoded_image_label.bind("<Double-Button-1>",
                                              lambda event: open_img(filename, cv2.imread(filename)))
                self.encoded_image_label.grid(row=14, column=1)
            elif filename.endswith('.gif'):
                gis.gif_save(self.encoded, filename)
                # show gif
                self.encoded_gif = Image.open(filename)
                self.encoded_frames = [tk.PhotoImage(file=filename, format='gif -index %i' % i) for i in
                                       range(self.encoded_gif.n_frames)]
                self.encoded_gif_label = tk.Label(self, image=self.encoded_frames[0], text="", height=100, width=100)
                self.encoded_gif_label.grid(row=14, column=1)
                self.encoded_gif_label.bind("<Double-Button-1>", lambda event: GifWindow(self, filename))
                self.encoded_gif_after_id = self.after(0,
                                                       lambda: play_gif(0, self.encoded_frames, self.encoded_gif_label,
                                                                        self.encoded_gif.info['duration'], self))

    def save_text(self):
        if self.checkbox.get() == 1:
            filename = asksaveasfilename(defaultextension=".txt",
                                         filetypes=[("Text File", ".txt"), ("All Files", "*.*")],
                                         initialfile="secret_message.txt")
            if filename:
                with open(filename, "w") as text:
                    text.write(self.secret_message.get('1.0', 'end-1c'))


class DocumentPage(customtkinter.CTkFrame):
    def __init__(self, master, controller, **kwargs):
        super().__init__(master, **kwargs)

        self.grid_columnconfigure((0, 1), weight=1)

        self.coverPath = None

        self.back_button = customtkinter.CTkButton(self, text="Back", command=lambda: controller.show_frame(MainMenu))
        self.back_button.grid(row=0, column=0, columnspan=2, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Document Steganography",
                                            font=customtkinter.CTkFont(size=20, weight="bold"))
        self.label.grid(row=1, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Settings For Docx File", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=2, column=0, padx=20, pady=10, columnspan=2)
        self.label = customtkinter.CTkLabel(self, text="Set Number of Bits", font=customtkinter.CTkFont(size=15))
        self.label.grid(row=3, column=0, padx=20, pady=2)
        self.bits_option_menu = customtkinter.CTkOptionMenu(self, values=["1", "2", "3", "4", "5"])
        self.bits_option_menu.grid(row=4, column=0, padx=20, pady=(10, 10))

        # switch_var = customtkinter.StringVar(value = "on")
        # self.switch = customtkinter.CTkSwitch(self, text = "Vanishing Steganograpy (Off)\nLSB Steganography (On)", variable = switch_var, onvalue = "on", offvalue="off")
        # self.switch.grid(row = 4, column = 1, padx = 20, pady = 10)

        self.label = customtkinter.CTkLabel(self, text="Cover", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=5, column=0, padx=20, pady=10, columnspan=2)
        self.upload_button = customtkinter.CTkButton(self, text="Click here or Drop an Document here to upload",
                                                     command=self.cover_on_change)
        self.upload_button.grid(row=6, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_button.drop_target_register(DND_FILES)
        self.upload_button.dnd_bind("<<Drop>>", self.cover_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Secret Message", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=7, column=0, padx=20, pady=10, columnspan=2)
        self.secret_message = customtkinter.CTkTextbox(self, height=100, width=500)
        self.secret_message.grid(row=8, column=0, sticky='w', columnspan=2)

        self.upload_txt_button = customtkinter.CTkButton(self,
                                                         text="Click here or Drop an txt here to upload a secret message",
                                                         command=self.txt_on_change)
        self.upload_txt_button.grid(row=9, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_txt_button.drop_target_register(DND_FILES)
        self.upload_txt_button.dnd_bind("<<Drop>>", self.txt_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Key", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=10, column=0, padx=20, pady=10, columnspan=2)
        self.key_message = customtkinter.CTkTextbox(self, height=40, width=500)
        self.key_message.grid(row=11, column=0, sticky='w', pady=10, columnspan=2)

        self.encode_button = customtkinter.CTkButton(self, text="Encode Document", command=self.encode_document)
        self.encode_button.grid(row=12, column=0, padx=20, pady=10)

        self.decode_button = customtkinter.CTkButton(self, text="Decode Document", command=self.decode_document)
        self.decode_button.grid(row=12, column=1, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Document Preview", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=14, column=0, padx=20, pady=10, columnspan=2)
        self.before_text = customtkinter.CTkTextbox(self, height=75, width=500)
        self.before_text.grid(row=15, column=0, sticky='w', columnspan=2)

    def txt_on_drop(self, event):
        self.txtPath = get_drop(event, ('', '.txt'))
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def txt_on_change(self):
        self.txtPath = get_path([('', '*.txt')])
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def cover_on_drop(self, event):
        self.coverPath = get_drop(event, ('.txt', '.xlsx', '.docx'))
        if self.coverPath is not None:
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=13, column=0)
            content = ""
            # Read the content of the file
            _, ext = os.path.splitext(self.coverPath)
            if ext.lower() == '.docx':
                doc = docx.Document(self.coverPath)
                content = [p.text for p in doc.paragraphs]
            elif ext.lower() == '.txt':
                with open(self.coverPath, encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            # Set the content in the "Before Encoding" textbox
            clear_text(self.secret_message)
            clear_text(self.before_text)
            self.before_text.insert('1.0', content)

    def cover_on_change(self):
        self.coverPath = get_path([('', '*.txt'), ('', '*.xlsx'), ('', '*.docx')])
        if self.coverPath is not None:
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=13, column=0)

            # Read the content of the file
            _, ext = os.path.splitext(self.coverPath)
            content = ""
            if ext.lower() == '.docx':
                doc = docx.Document(self.coverPath)
                content = [p.text for p in doc.paragraphs]
            elif ext.lower() == '.txt':
                with open(self.coverPath, encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            # Set the content in the "Before Encoding" textbox
            clear_text(self.secret_message)
            clear_text(self.before_text)
            self.before_text.insert('1.0', content)

    def encode_document(self):
        # Error Handling if the User has not select a Cover File
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return
        _, ext = os.path.splitext(self.coverPath)

        if len(self.secret_message.get('1.0', 'end-1c')) == 0:
            messagebox.showerror("Error", "Please enter a secret message.")
            return

        self.payloadText = self.secret_message.get('1.0', 'end-1c')
        message = self.key_message.get('1.0', 'end-1c')
        if ext == ".docx":
            try:
                if (len(message) > 0):
                    if (enc.has_repeating_characters(message)):
                        messagebox.showerror("Error",
                                             "Secret key has repeating characters, please use a unique string sequence")
                        return
                    else:
                        self.payloadText = enc.encryptMessage(self.payloadText, message)
                self.encoded = supported_types[ext.lower()][0](self.coverPath, self.payloadText,
                                                               int(self.bits_option_menu.get()))
            except ValueError as e:
                messagebox.showerror("Error", str(e))
                return
            self.save_as(ext)
        elif ext == ".txt":
            # Error Handling if the Secret can fit into the Cover File
            max_no_of_bytes = dms.max_number_of_bytes(self.coverPath)
            if dms.check_secret_can_fit(max_no_of_bytes, self.payloadText) == False:
                messagebox.showerror("Error", "Insufficient bytes, need a bigger text file or reduce payload")
                return
            self.save_as(ext, secret=self.payloadText)
        elif ext == ".xlsx":
            try:
                if (len(message) > 0):
                    if (enc.has_repeating_characters(message)):
                        messagebox.showerror("Error",
                                             "Secret key has repeating characters, please use a unique string sequence")
                        return
                    else:
                        self.payloadText = enc.encryptMessage(self.payloadText, message)
                self.encoded = supported_types[ext.lower()][0](self.coverPath, self.payloadText,
                                                               int(self.bits_option_menu.get()))
            except ValueError as e:
                messagebox.showerror("Error", str(e))
                return
            self.save_as(ext)

    def decode_document(self):
        # Error Handling if the User has not select a Cover File
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return

        _, ext = os.path.splitext(self.coverPath)
        message = self.key_message.get('1.0', 'end-1c')
        if ext == ".docx":
            text = wd_lsb_s.decode(self.coverPath, int(self.bits_option_menu.get()))
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                text = enc.decryptMessage(text, self.key_message.get('1.0', 'end-1c'))
            if text == 0:
                messagebox.showerror("Error", "Wrong Key")
                return
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', text)
        elif ext == ".txt":
            text = dms.decode(self.coverPath)
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                text = enc.decryptMessage(text, self.key_message.get('1.0', 'end-1c'))
                if text == 0:
                    messagebox.showerror("Error", "Wrong Key")
                    return
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', text)
        elif ext == ".xlsx":
            try:
                text = xls.decode(self.coverPath, int(self.bits_option_menu.get()))
            except ValueError as e:
                messagebox.showerror("Error", str(e))
                return
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                text = enc.decryptMessage(text, self.key_message.get('1.0', 'end-1c'))
                if text == 0:
                    messagebox.showerror("Error", "Wrong Key")
                    return
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', text)

    def save_as(self, ext, secret=None):
        # save as prompt
        filename = asksaveasfilename(defaultextension=ext, filetypes=[("Same as original", ext), ("All Files", "*.*")])
        if filename:
            if filename.endswith(('.txt')):
                dms.encode(self.coverPath, filename, secret)
                print("\nEncoded the data successfully in the document file")
            elif filename.endswith(('.docx')):
                wd_lsb_s.save(filename, self.encoded)
            elif filename.endswith(('.xlsx')):
                xls.save(filename, self.encoded)


class AudioPage(customtkinter.CTkFrame):
    def __init__(self, master, controller, **kwargs):
        super().__init__(master, **kwargs)

        self.grid_columnconfigure((0, 1), weight=1)
        self.coverPath = None
        self.audio = None

        self.back_button = customtkinter.CTkButton(self, text="Back", command=lambda: controller.show_frame(MainMenu))
        self.back_button.grid(row=0, column=0, columnspan=2, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Audio Steganography",
                                            font=customtkinter.CTkFont(size=20, weight="bold"))
        self.label.grid(row=1, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Cover", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=2, column=0, padx=20, pady=10, columnspan=2)
        self.upload_button = customtkinter.CTkButton(self, text="Click here or Drop an Audio here to upload",
                                                     command=self.cover_on_change)
        self.upload_button.grid(row=3, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_button.drop_target_register(DND_FILES)
        self.upload_button.dnd_bind("<<Drop>>", self.cover_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Secret Message", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=4, column=0, padx=20, pady=10, columnspan=2)
        self.secret_message = customtkinter.CTkTextbox(self, height=100, width=500)
        self.secret_message.grid(row=5, column=0, sticky='w', columnspan=2)

        self.upload_txt_button = customtkinter.CTkButton(self,
                                                         text="Click here or Drop an txt here to upload a secret message",
                                                         command=self.txt_on_change)
        self.upload_txt_button.grid(row=6, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_txt_button.drop_target_register(DND_FILES)
        self.upload_txt_button.dnd_bind("<<Drop>>", self.txt_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Key", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=7, column=0, padx=20, pady=10, columnspan=2)
        self.key_message = customtkinter.CTkTextbox(self, height=40, width=500)
        self.key_message.grid(row=8, column=0, sticky='w', pady=10, columnspan=2)

        self.encode_button = customtkinter.CTkButton(self, text="Encode Audio", command=self.encode_audio)
        self.encode_button.grid(row=9, column=0, padx=20, pady=10)

        self.decode_button = customtkinter.CTkButton(self, text="Decode Audio", command=self.decode_audio)
        self.decode_button.grid(row=9, column=1, padx=20, pady=10)

        self.play_button = customtkinter.CTkButton(self, text="Play Audio",
                                                   font=customtkinter.CTkFont(size=20, weight="bold"),
                                                   command=self.play_audio)
        self.play_button.grid(row=10, column=0, padx=20, pady=10, columnspan=1)

        self.stop_button = customtkinter.CTkButton(self, text="Stop Audio",
                                                   font=customtkinter.CTkFont(size=20, weight="bold"),
                                                   command=self.stop_audio)
        self.stop_button.grid(row=10, column=1, padx=20, pady=10, columnspan=1)

    def txt_on_drop(self, event):
        self.txtPath = get_drop(event, ('', '.txt'))
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def txt_on_change(self):
        self.txtPath = get_path([('', '*.txt')])
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def play_audio(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "No audio file is uploaded")
            return
        self.audio = vlc.MediaPlayer(self.coverPath)
        self.audio.play()

    def stop_audio(self):
        if self.audio is not None:
            self.audio.stop()

    def cover_on_drop(self, event):
        self.coverPath = get_drop(event, ('.wav', '.mp3'))
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=11, column=0)

    def cover_on_change(self):
        self.coverPath = get_path([('', '*.wav'), ('', '*.mp3')])
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=11, column=0)

    def encode_audio(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return
        _, ext = os.path.splitext(self.coverPath)

        if len(self.secret_message.get('1.0', 'end-1c')) == 0:
            messagebox.showerror("Error", "Please enter a secret message.")
            return

        self.payloadText = self.secret_message.get('1.0', 'end-1c')
        message = self.key_message.get('1.0', 'end-1c')
        if (len(message) > 0):
            if (enc.has_repeating_characters(message)):
                messagebox.showerror("Error",
                                     "Secret key has repeating characters, please use a unique string sequence")
                return
            else:
                self.payloadText = enc.encryptMessage(self.payloadText, message)
        self.encoded = supported_types[ext.lower()][0](self.coverPath, self.payloadText)
        self.save_as(ext)

    def decode_audio(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return
        try:
            text = auds.decode_aud_data(self.coverPath)
            message = self.key_message.get('1.0', 'end-1c')
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                else:
                    text = enc.decryptMessage(text, self.key_message.get('1.0', 'end-1c'))
            if text == 0:
                messagebox.showerror("Error", "Wrong Key")
                return
        except Exception:
            messagebox.showerror("Error", "Audio file contains no signs of steganography")
            return
        _, ext = os.path.splitext(self.coverPath)
        if ext == ".mp3":
            auds.clean_tmp()
        clear_text(self.secret_message)
        self.secret_message.insert('1.0', text)

    def save_as(self, ext, secret=None):
        # save as prompt
        filename = asksaveasfilename(defaultextension=ext, filetypes=[("Same as original", ext), ("All Files", "*.*")])
        if filename:
            if filename.endswith(('.wav')) and ext != ".mp3":
                song = wave.open(self.coverPath, mode='rb')
                with wave.open(filename, 'wb') as fd:
                    fd.setparams(song.getparams())
                    fd.writeframes(self.encoded)
                    print("\nEncoded the data successfully in the audio file")
            elif ext == ".mp3":
                _, ext = os.path.splitext(filename)
                song_params = auds.get_song_parameters(self.coverPath)
                with wave.open(filename[:-3] + "wav", 'wb') as fd:
                    fd.setparams(song_params)
                    fd.writeframes(self.encoded)
                    print("\nEncoded the data successfully in the audio file")
                auds.convertWAVToMP3(filename[:-3] + "wav")
                # Remove temporary WAV file
                os.remove(filename[:-3] + "wav")
                auds.clean_tmp()


class VideoPage(customtkinter.CTkFrame):
    def __init__(self, master, controller, **kwargs):
        super().__init__(master, **kwargs)

        self.grid_columnconfigure((0, 1), weight=1)

        self.coverPath = None
        self.encoded_video_name = None
        self.cover_player = None
        self.encoded_player = None

        self.back_button = customtkinter.CTkButton(self, text="Back", command=lambda: controller.show_frame(MainMenu))
        self.back_button.grid(row=0, column=0, columnspan=2, padx=20, pady=10)

        self.label = customtkinter.CTkLabel(self, text="Video Steganography",
                                            font=customtkinter.CTkFont(size=20, weight="bold"))
        self.label.grid(row=1, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Settings", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=2, column=0, padx=20, pady=10, columnspan=2)

        self.label = customtkinter.CTkLabel(self, text="Set Frame Number to Encode Secret in",
                                            font=customtkinter.CTkFont(size=15))
        self.label.grid(row=3, column=0, padx=5, pady=2)
        self.frame_option = customtkinter.CTkEntry(self, placeholder_text="Frame Number")
        self.frame_option.grid(row=4, column=0, padx=5, pady=(10, 10))

        self.label = customtkinter.CTkLabel(self, text="Set Number of Bits", font=customtkinter.CTkFont(size=15))
        self.label.grid(row=3, column=1, padx=5, pady=2)
        self.bits_option_menu = customtkinter.CTkOptionMenu(self, values=["1", "2", "3", "4", "5"])
        self.bits_option_menu.grid(row=4, column=1, padx=5, pady=(10, 10))

        self.label = customtkinter.CTkLabel(self, text="Cover", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=5, column=0, padx=20, pady=10, columnspan=2)
        self.upload_button = customtkinter.CTkButton(self, text="Click here or Drop a Video here to upload",
                                                     command=self.cover_on_change)
        self.upload_button.grid(row=6, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_button.drop_target_register(DND_FILES)
        self.upload_button.dnd_bind("<<Drop>>", self.cover_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Secret Message", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=7, column=0, padx=20, pady=10, columnspan=2)
        self.secret_message = customtkinter.CTkTextbox(self, height=100, width=500)
        self.secret_message.grid(row=8, column=0, sticky='w', columnspan=2)

        self.upload_txt_button = customtkinter.CTkButton(self,
                                                         text="Click here or Drop an txt here to upload a secret message",
                                                         command=self.txt_on_change)
        self.upload_txt_button.grid(row=9, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        self.upload_txt_button.drop_target_register(DND_FILES)
        self.upload_txt_button.dnd_bind("<<Drop>>", self.txt_on_drop)

        self.label = customtkinter.CTkLabel(self, text="Key", font=customtkinter.CTkFont(size=20))
        self.label.grid(row=10, column=0, padx=20, pady=10, columnspan=2)
        self.key_message = customtkinter.CTkTextbox(self, height=40, width=500)
        self.key_message.grid(row=11, column=0, sticky='w', pady=10, columnspan=2)

        self.encode_button = customtkinter.CTkButton(self, text="Encode Video", command=self.encode_video)
        self.encode_button.grid(row=12, column=0, padx=20, pady=10)

        self.decode_button = customtkinter.CTkButton(self, text="Decode Video", command=self.decode_video)
        self.decode_button.grid(row=12, column=1, padx=20, pady=10)

        self.cover_video = customtkinter.CTkButton(self, text="Play Cover Video", command=self.play_cover_video)
        self.cover_video.grid(row=13, column=0, padx=20, pady=10)

        self.encoded_video = customtkinter.CTkButton(self, text="Play Encoded Video", command=self.play_encoded_video)
        self.encoded_video.grid(row=13, column=1, padx=20, pady=10)

        self.cover_video = customtkinter.CTkButton(self, text="Stop Cover Video", command=self.stop_cover_video)
        self.cover_video.grid(row=14, column=0, padx=20, pady=10)

        self.cover_video = customtkinter.CTkButton(self, text="Stop Encoded Video", command=self.stop_encoded_video)
        self.cover_video.grid(row=14, column=1, padx=20, pady=10)

    def txt_on_drop(self, event):
        self.txtPath = get_drop(event, ('', '.txt'))
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def txt_on_change(self):
        self.txtPath = get_path([('', '*.txt')])
        if self.txtPath is not None:
            with open(self.txtPath) as f:
                contents = f.read()
            clear_text(self.secret_message)
            self.secret_message.insert('1.0', contents)

    def play_cover_video(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "No Cover Video Added")
            return
        self.vlc_instance_cover = vlc.Instance()
        self.cover_player = self.vlc_instance_cover.media_player_new()
        self.cover_media = self.vlc_instance_cover.media_new(self.coverPath)
        self.cover_player.set_media(self.cover_media)
        self.cover_player.play()

    def play_encoded_video(self):
        if self.encoded_video_name is None:
            messagebox.showerror("Error", "No Encoded Video Added")
            return
        self.vlc_instance_encoded = vlc.Instance()
        self.encoded_player = self.vlc_instance_encoded.media_player_new()
        self.encoded_media = self.vlc_instance_encoded.media_new(self.encoded_video_name)
        self.encoded_player.set_media(self.encoded_media)
        self.encoded_player.play()

    def stop_cover_video(self):
        if self.coverPath is None or self.cover_player is None:
            messagebox.showerror("Error", "No Cover Video Added")
            return
        self.cover_player.release()

    def stop_encoded_video(self):
        if self.encoded_video_name is None or self.encoded_player is None:
            messagebox.showerror("Error", "No Encoded Video Added")
            return
        self.encoded_player.release()

    def cover_on_drop(self, event):
        self.coverPath = get_drop(event, ('', '.mp4'))
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=15, column=0)

    def cover_on_change(self):
        self.coverPath = get_path([('', '*.mp4')])
        if self.coverPath is not None:
            clear_text(self.secret_message)
            self.success_label = customtkinter.CTkLabel(self, text="Successfully uploaded file")
            self.success_label.grid(row=15, column=0)

    def encode_video(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return

        _, ext = os.path.splitext(self.coverPath)
        if (len(self.secret_message.get('1.0', 'end-1c')) == 0):
            messagebox.showerror("Error", "Please enter a secret message")
            return

        self.payloadText = self.secret_message.get('1.0', 'end-1c')
        try:
            message = self.key_message.get('1.0', 'end-1c')
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                else:
                    self.payloadText = enc.encryptMessage(self.payloadText, message)
            supported_types[ext.lower()][0](self.coverPath, self.payloadText, self.frame_option.get(),
                                            self.bits_option_menu.get())
        except ValueError as e:
            messagebox.showerror("Error", str(e))
            vids.clean_tmp()
            return
        self.save_as(ext)

    def decode_video(self):
        if self.coverPath is None:
            messagebox.showerror("Error", "Please select or drop a cover item first!")
            return

        if len(self.frame_option.get()) == 0:
            messagebox.showerror("Error", "Please enter a frame number")
            return

        _, ext = os.path.splitext(self.coverPath)
        try:
            message = self.key_message.get('1.0', 'end-1c')
            text = supported_types[ext.lower()][1](self.coverPath, self.frame_option.get(), self.bits_option_menu.get())
            if (len(message) > 0):
                if (enc.has_repeating_characters(message)):
                    messagebox.showerror("Error",
                                         "Secret key has repeating characters, please use a unique string sequence")
                    return
                else:
                    text = enc.decryptMessage(text, self.key_message.get('1.0', 'end-1c'))
            if text == 0:
                messagebox.showerror("Error", "Wrong Key")
                return
        except ValueError as e:
            messagebox.showerror("Error", str(e))
            vids.clean_tmp()
            return
        clear_text(self.secret_message)
        self.secret_message.insert('1.0', text)

    def save_as(self, ext):
        filename = asksaveasfilename(defaultextension=ext, filetypes=[("Same as original", ext), ("All files", "*.*")])
        if filename:
            if filename.endswith(('.mp4')):
                vids.save_as(filename)
                self.encoded_video_name = filename


class App(customtkinter.CTk, TkinterDnD.DnDWrapper):
    def __init__(self):
        super().__init__()
        self.TkdndVersion = TkinterDnD._require(self)

        self.title("I am a stegosaurus")
        self.geometry("500x950")

        container = customtkinter.CTkFrame(master=self)
        container.pack(side="top", fill="both", expand=True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)
        self.frames = {}

        for F in (MainMenu, ImagePage, DocumentPage, AudioPage, VideoPage):
            frame = F(container, self)

            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(MainMenu)

    def show_frame(self, cont):
        frame = self.frames[cont]
        frame.tkraise()


app = App()
app.mainloop()
